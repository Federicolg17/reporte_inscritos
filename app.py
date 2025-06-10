import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import numpy as np
import tempfile
import io
from pathlib import Path

# Configuración de la página
st.set_page_config(
    page_title="Generador de Reportes de Inscripciones",
    page_icon="📊",
    layout="wide"
)

# Título principal
st.title("📊 Generador de Reportes de Inscripciones a Cursos")
st.markdown("---")

# Descripción
st.markdown("""
Esta aplicación genera automáticamente un reporte detallado de inscripciones a cursos a partir de un archivo Excel.
El reporte incluye:
- Resumen estadístico general
- Gráfico de distribución por curso
- Tablas detalladas de inscritos por curso
""")

# Sidebar para la carga de archivos
st.sidebar.header("📁 Carga de Datos")
uploaded_file = st.sidebar.file_uploader(
    "Selecciona el archivo Excel con los datos de inscripciones",
    type=['xlsx', 'xls'],
    help="El archivo debe contener las columnas: 'Nombre y apellidos completos', 'Hora de inicio', 'Curso de interés', 'Correo de contacto'"
)

def process_data(df):
    """Procesa los datos y genera las estadísticas necesarias"""
    # Configurar pandas para mostrar todas las columnas
    pd.set_option('display.max_columns', None)
    
    # 1. Cantidad de personas inscritas (valores únicos)
    personas_unicas = df['Nombre y apellidos completos'].nunique()
    
    # 2. Fecha máxima y mínima de inscripción
    df['Hora de inicio'] = pd.to_datetime(df['Hora de inicio'])
    fecha_min = df['Hora de inicio'].min()
    fecha_max = df['Hora de inicio'].max()
    
    # 3. Contar inscripciones por curso
    inscripciones_por_curso = df['Curso de interés'].value_counts().reset_index()
    inscripciones_por_curso.columns = ['Curso', 'Cantidad']
    
    return personas_unicas, fecha_min, fecha_max, inscripciones_por_curso

def create_chart(inscripciones_por_curso):
    """Crea el gráfico de barras de inscripciones por curso"""
    fig, ax = plt.subplots(figsize=(12, 8))
    
    bars = ax.bar(inscripciones_por_curso['Curso'], inscripciones_por_curso['Cantidad'], 
                  color='skyblue', edgecolor='navy', linewidth=1.2)
    
    # Añadir etiquetas en las barras
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                f'{int(height)}', ha='center', va='bottom', fontweight='bold')
    
    ax.set_xlabel('Curso', fontsize=12, fontweight='bold')
    ax.set_ylabel('Cantidad de Inscritos', fontsize=12, fontweight='bold')
    ax.set_title('Inscripciones por Curso', fontsize=14, fontweight='bold')
    ax.tick_params(axis='x', rotation=45)
    
    plt.tight_layout()
    return fig

def create_word_report(df, personas_unicas, fecha_min, fecha_max, inscripciones_por_curso, chart_path):
    """Crea el documento Word con el reporte completo"""
    documento = Document()
    
    # Configurar estilo de título
    titulo = documento.add_heading('REPORTE DE INSCRIPCIONES A CURSOS', level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir fecha de elaboración del informe
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    parrafo_fecha = documento.add_paragraph(f'Fecha de elaboración: {fecha_actual}')
    parrafo_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Añadir elaborado por
    parrafo_autor = documento.add_paragraph('Elaborado por: Sistema Automatizado de Reportes')
    parrafo_autor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    documento.add_heading('Resumen General', level=1)
    
    # Añadir información de cantidad de inscritos
    documento.add_paragraph(f'Total de personas inscritas (valores únicos): {personas_unicas}')
    
    # Añadir información de fechas
    documento.add_paragraph(f'Fecha de inicio de inscripciones: {fecha_min.strftime("%d/%m/%Y %H:%M:%S")}')
    documento.add_paragraph(f'Fecha de finalización de inscripciones: {fecha_max.strftime("%d/%m/%Y %H:%M:%S")}')
    
    # Añadir gráfico
    documento.add_heading('Distribución de Inscripciones por Curso', level=1)
    documento.add_picture(chart_path, width=Inches(6.0))
    documento.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir tablas por curso
    documento.add_heading('Detalle de Inscritos por Curso', level=1)
    
    # Obtener lista de cursos únicos
    cursos_unicos = df['Curso de interés'].unique()
    
    for curso in cursos_unicos:
        # Filtrar datos para este curso
        df_curso = df[df['Curso de interés'] == curso]
        
        # Seleccionar columnas relevantes y eliminar duplicados por nombre
        df_curso = df_curso[['Nombre y apellidos completos', 'Correo de contacto']].drop_duplicates(subset=['Nombre y apellidos completos'])
        
        # Añadir subtítulo para el curso
        documento.add_heading(f'Curso: {curso}', level=2)
        
        # Añadir cantidad de inscritos
        documento.add_paragraph(f'Total de inscritos: {len(df_curso)}')
        
        # Crear tabla
        tabla = documento.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        
        # Establecer encabezados
        encabezados = tabla.rows[0].cells
        encabezados[0].text = 'Nombre y Apellidos'
        encabezados[1].text = 'Correo de Contacto'
        
        # Añadir datos a la tabla
        for _, fila in df_curso.iterrows():
            celdas = tabla.add_row().cells
            celdas[0].text = str(fila['Nombre y apellidos completos'])
            celdas[1].text = str(fila['Correo de contacto'])
        
        # Añadir espacio después de cada tabla
        documento.add_paragraph('')
    
    return documento

# Procesamiento principal
if uploaded_file is not None:
    try:
        # Leer el archivo Excel
        df = pd.read_excel(uploaded_file)
        
        # Validar columnas requeridas
        required_columns = ['Nombre y apellidos completos', 'Hora de inicio', 'Curso de interés', 'Correo de contacto']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            st.error(f"❌ Faltan las siguientes columnas en el archivo: {', '.join(missing_columns)}")
            st.info("Asegúrate de que tu archivo Excel contenga todas las columnas requeridas.")
        else:
            # Mostrar información básica del archivo
            st.success("✅ Archivo cargado exitosamente!")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Total de registros", len(df))
            with col2:
                st.metric("Columnas detectadas", len(df.columns))
            with col3:
                st.metric("Cursos únicos", df['Curso de interés'].nunique())
            
            # Procesar datos
            personas_unicas, fecha_min, fecha_max, inscripciones_por_curso = process_data(df)
            
            # Mostrar resumen estadístico
            st.markdown("---")
            st.header("📈 Resumen Estadístico")
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Personas únicas inscritas", personas_unicas)
                st.metric("Fecha de primera inscripción", fecha_min.strftime("%d/%m/%Y"))
            with col2:
                st.metric("Total de cursos", len(inscripciones_por_curso))
                st.metric("Fecha de última inscripción", fecha_max.strftime("%d/%m/%Y"))
            
            # Mostrar gráfico
            st.markdown("---")
            st.header("📊 Distribución de Inscripciones por Curso")
            
            fig = create_chart(inscripciones_por_curso)
            st.pyplot(fig)
            
            # Mostrar tabla de inscripciones por curso
            st.markdown("---")
            st.header("📋 Tabla de Inscripciones por Curso")
            st.dataframe(inscripciones_por_curso, use_container_width=True)
            
            # Generar reporte Word
            st.markdown("---")
            st.header("📄 Generar Reporte Completo")
            
            if st.button("🔄 Generar Reporte Word", type="primary"):
                with st.spinner("Generando reporte..."):
                    # Crear archivos temporales
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_chart:
                        fig.savefig(tmp_chart.name, dpi=300, bbox_inches='tight')
                        chart_path = tmp_chart.name
                    
                    # Crear documento Word
                    documento = create_word_report(df, personas_unicas, fecha_min, fecha_max, 
                                                 inscripciones_por_curso, chart_path)
                    
                    # Guardar documento en memoria
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc:
                        documento.save(tmp_doc.name)
                        
                        # Leer el archivo para descarga
                        with open(tmp_doc.name, 'rb') as file:
                            doc_bytes = file.read()
                    
                    # Limpiar archivos temporales
                    os.unlink(chart_path)
                    os.unlink(tmp_doc.name)
                
                st.success("✅ Reporte generado exitosamente!")
                
                # Botón de descarga
                st.download_button(
                    label="⬇️ Descargar Reporte Word",
                    data=doc_bytes,
                    file_name=f"reporte_inscripciones_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Mostrar vista previa de datos
            st.markdown("---")
            st.header("👁️ Vista Previa de Datos")
            st.dataframe(df.head(10), use_container_width=True)
            
    except Exception as e:
        st.error(f"❌ Error al procesar el archivo: {str(e)}")
        st.info("Verifica que el archivo Excel esté en el formato correcto y no esté corrupto.")

else:
    # Mostrar instrucciones cuando no hay archivo cargado
    st.info("👆 Por favor, carga un archivo Excel para comenzar.")
    
    # Mostrar ejemplo de estructura de datos
    st.markdown("---")
    st.header("📋 Estructura de Datos Requerida")
    st.markdown("Tu archivo Excel debe contener las siguientes columnas:")
    
    example_data = {
        'Nombre y apellidos completos': ['Juan Pérez García', 'María López Rodríguez', 'Carlos Martín Sánchez'],
        'Hora de inicio': ['2024-01-15 09:30:00', '2024-01-16 14:20:00', '2024-01-17 11:45:00'],
        'Curso de interés': ['Python Básico', 'Excel Avanzado', 'Python Básico'],
        'Correo de contacto': ['juan.perez@email.com', 'maria.lopez@email.com', 'carlos.martin@email.com']
    }
    
    example_df = pd.DataFrame(example_data)
    st.dataframe(example_df, use_container_width=True)

# Footer
st.markdown("---")
st.markdown("Desarrollado con ❤️ usando Streamlit | © 2025")
